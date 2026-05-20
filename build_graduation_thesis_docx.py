# -*- coding: utf-8 -*-
"""生成毕业设计论文 Word 文档（需本地安装 python-docx）。运行: python build_graduation_thesis_docx.py"""
from __future__ import annotations

import os
import shutil
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

import graduation_thesis_expansion as _thesis_exp


def set_run_eastasia_font(run, font_name: str = "宋体", size_pt: float = 12) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    align=None,
    first_line_indent_cm: float | None = 0.74,
    font_name: str = "宋体",
    size_pt: float = 12,
    line_spacing: float | None = 1.5,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if line_spacing is not None:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.bold = bold
    set_run_eastasia_font(run, font_name, size_pt)


def add_heading_cn(doc: Document, text: str, level: int) -> None:
    # 使用 Word 内置标题样式以便「引用 → 目录」自动生成
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        set_run_eastasia_font(run, "黑体" if level == 1 else "黑体", 16 if level == 1 else (14 if level == 2 else 12))
        run.bold = True


def add_caption(doc: Document, text: str, *, kind: str = "table") -> None:
    """插入题注风格段落：用于表格题注/图题题注。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    set_run_eastasia_font(run, "宋体", 12)


def add_title_block(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("（学校、学院、专业、学号、姓名、指导教师等信息请按学校模板替换）\n\n")
    set_run_eastasia_font(r, "宋体", 10.5)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("本科毕业设计（论文）")
    set_run_eastasia_font(r2, "黑体", 18)
    r2.bold = True
    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("基于语言大模型的水稻分子机理数据检索和知识问答")
    set_run_eastasia_font(r3, "黑体", 16)
    r3.bold = True
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(str(date.today().year) + " 年 5 月")
    set_run_eastasia_font(rr, "宋体", 12)
    doc.add_page_break()


def document_to_plain_text(doc: Document) -> str:
    lines: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    return "\n\n".join(lines)


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    add_title_block(doc)

    # 中文摘要
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = h.add_run("摘  要")
    set_run_eastasia_font(hr, "黑体", 16)
    hr.bold = True
    doc.add_paragraph()

    abstract_zh = (
        "水稻分子机理研究高度依赖 Gene Ontology（GO）等结构化知识。将自然语言中的机理表述可靠地映射为"
        "GO 标识符（实体识别 / 术语归一）是后续检索与问答质量的关键环节；仅靠关键词或纯大模型生成易出现"
        "漏检、误检与幻觉。本文以「GO 术语实体识别」为研究与工程上的工作重点，在完成数据检索与"
        "检索增强生成（RAG）问答的基础上，系统实现并评测多类识别管线及其组合策略。主要工作包括："
        "（1）实现不少于八类可区分的 GO 识别路线：词典匹配、NLTK 句法级匹配、向量语义检索、纯大模型抽取"
        "与归一化、检索增强式大模型判别（先召回候选再约束输出）、NLTK 与大模型级联、分词—候选—大模型"
        "流水线，以及字典—向量—大模型加权融合的 Ensemble 集成；集成策略提供 strict / balanced / recall"
        "三组参数预设，用于在精度、召回与误报之间折中。（2）构建向量索引与语义检索子系统，支撑候选召回、"
        "RAG 证据链及识别模块中的向量分支。（3）实现 RAG 问答与可选的「识别结果驱动的查询扩展」，形成"
        "识别—检索—生成闭环。在自建 GO 标注基准（含中英及混合句）及离线 benchmark 脚本输出上进行"
        "横向与纵向对比；综合 Precision、Recall、F1、负例误报与耗时等指标。实验表明：保守约束的大模型"
        "识别在综合 F1 与误报控制上优于若干传统基线；检索增强式识别有利于约束输出空间；不同集成预设"
        "对召回与误报影响显著。本文工作为面向水稻机理文本的 GO 落地应用提供了可复现的方法谱系与实验依据。"
    )
    add_para(doc, abstract_zh, first_line_indent_cm=0.74)

    doc.add_paragraph()
    kw = doc.add_paragraph()
    kwr = kw.add_run(
        "关键词：大语言模型；Gene Ontology；实体识别；检索增强生成；水稻；知识问答；向量检索"
    )
    set_run_eastasia_font(kwr, "宋体", 12)
    kwr.bold = True

    doc.add_page_break()

    # 英文摘要
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2r = h2.add_run("Abstract")
    set_run_eastasia_font(h2r, "Times New Roman", 16)
    h2r.bold = True
    doc.add_paragraph()

    abstract_en = (
        "Rice molecular mechanism research relies on structured resources such as Gene Ontology (GO). Reliably "
        "mapping free text to GO identifiers (entity recognition / term normalization) is critical for downstream "
        "retrieval and question answering; keyword-only or unconstrained LLM generation suffers from missed "
        "mentions, false positives, and hallucinations. This thesis emphasizes GO term recognition as the primary "
        "research and engineering focus, while still delivering semantic retrieval and retrieval-augmented generation "
        "(RAG) QA. The contributions include: (1) implementing and benchmarking eight distinguishable recognition "
        "pipelines—dictionary matching, NLTK-style syntactic matching, dense vector retrieval, LLM extraction with "
        "normalization, retrieval-augmented LLM judging over a candidate list, NLTK–LLM and token-candidate–LLM "
        "cascades, and a Dict–Vector–LLM ensemble with three named presets (strict, balanced, recall) to trade off "
        "precision, recall, and false alarms; (2) building FAISS-backed vector indexes to support candidate recall, "
        "RAG evidence, and the vector branch of recognition; (3) integrating optional query expansion driven by "
        "recognition outputs to close the recognize–retrieve–generate loop. Experiments use curated GO-labeled "
        "benchmarks (English, Chinese, and mixed) and offline benchmark scripts, reporting Precision, Recall, F1, "
        "negative false-positive rates, and latency. Results show that conservatively constrained LLM recognition "
        "improves the F1–false-positive trade-off versus several classical baselines; retrieval-augmented judging "
        "helps constrain the output space; ensemble presets strongly affect recall and noise. The work provides a "
        "reproducible method spectrum and empirical evidence for GO-centric rice mechanism text applications."
    )
    pe = doc.add_paragraph()
    pe.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pe.paragraph_format.line_spacing = 1.5
    pe.paragraph_format.first_line_indent = Cm(0.74)
    rer = pe.add_run(abstract_en)
    rer.font.name = "Times New Roman"
    rer.font.size = Pt(12)

    doc.add_paragraph()
    kwe = doc.add_paragraph()
    kwer = kwe.add_run(
        "Keywords: large language model; Gene Ontology; named entity recognition; retrieval-augmented generation; "
        "rice; knowledge question answering; vector retrieval"
    )
    kwer.font.name = "Times New Roman"
    kwer.font.size = Pt(12)
    kwer.bold = True

    doc.add_page_break()

    # 目录说明
    toc = doc.add_paragraph()
    tocr = toc.add_run("目  录")
    set_run_eastasia_font(tocr, "黑体", 16)
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_para(
        doc,
        "本文各级标题已应用 Word「标题 1 / 标题 2 / 标题 3」样式。请在 Microsoft Word 中选择："
        "「引用」→「目录」→「自动目录」，即可生成可更新的目录（生成后按 Ctrl+A 全选再按 F9 可更新页码）。",
        first_line_indent_cm=0.74,
    )
    doc.add_page_break()

    # ========== 第一章 ==========
    add_heading_cn(doc, "第一章 绪论", 1)

    add_heading_cn(doc, "1.1 研究背景与意义", 2)
    add_para(
        doc,
        "水稻是重要的粮食作物，其产量与品质形成受到众多分子过程协同调控。随着高通量测序与功能基因组学"
        "的发展，大量与生长发育、逆境响应及代谢调控相关的基因与通路被不断报道。Gene Ontology（GO）"
        "通过分子功能（Molecular Function）、生物过程（Biological Process）与细胞组分（Cellular Component）"
        "三个命名空间，为基因产物功能提供了可计算、可共享的标准化描述框架，已成为分子机理阐释与知识整合"
        "的重要基础设施。",
    )
    add_para(
        doc,
        "然而，科研人员在阅读文献与整理假说时，往往以自然语言提出问题：句中究竟对应哪些 GO 术语、"
        "非标准表述如何归一到权威 GO:ID，直接决定后续检索与生成是否「对题」。传统检索多依赖关键词匹配，"
        "难以处理同义表述、中英混合及语义近邻但生物学含义不同等情况；纯大模型端到端抽取则易出现幻觉"
        "GO 或负例误报。因而，围绕 GO 的实体识别（含多策略融合与可评测基准）应作为系统能力与论文叙述的"
        "核心之一；在此之上再叠加向量检索与 RAG 问答，才能形成可解释、可对照的完整链路。",
    )
    add_para(
        doc,
        "从国家粮食安全与分子设计育种战略看，将分散在文献与数据库中的机理知识快速沉淀为可计算标签，"
        "有助于加速基因功能验证与性状聚合分析。GO 作为跨物种共享的本体，为水稻与其他模式生物的"
        "比较基因组学研究提供了对齐坐标；因而面向 GO 的自然语言接口不仅服务于检索效率，也服务于"
        "知识发现的长尾需求。",
    )
    add_para(
        doc,
        "从计算机科学与技术专业视角看，本课题综合了信息检索、自然语言处理、知识工程与分布式服务"
        "等核心课程知识点：既有经典倒排与 n-gram 思想，也有深度向量与大模型推理；既要求算法正确性，"
        "也要求系统可部署与实验可复现，符合本科毕业设计对综合能力检验的定位。",
    )

    add_heading_cn(doc, "1.2 国内外研究现状", 2)
    add_para(
        doc,
        "在生物医学与自然语言处理交叉领域，知识图谱问答、检索增强生成（RAG）以及本体术语识别（NER / "
        "grounding）已形成较为活跃的研究方向。知识图谱问答通常围绕结构化查询或语义解析展开；RAG 通过"
        "将外部语料或知识库检索结果作为条件上下文，显著缓解纯生成模型的幻觉问题。针对 GO 等大规模本体，"
        "向量检索与词典匹配相结合的术语归一化方法在注释与文献挖掘任务中得到广泛应用。",
    )
    add_para(
        doc,
        "现有不足可概括为：通用开放域问答难以直接迁移到 GO 这类大规模受控词表；单一词典或单一向量通路"
        "往往在召回与精度上顾此失彼；大模型若无候选约束则输出空间过大、难以保证与本体一致。本文在综述"
        "基础上，面向水稻机理文本给出可落地的多路线识别体系，并在统一基准上系统对比「单路—级联—"
        "检索增强—加权集成」及多组超参预设，再与检索、RAG 问答模块联评。",
    )

    _thesis_exp.expand_chapter1(doc, add_para, add_heading_cn)

    add_heading_cn(doc, "1.3 研究目标", 2)
    add_para(
        doc,
        "本文研究目标为：在 GO 约束下，（1）建立覆盖多类范式的 GO 术语识别子系统，并支持集成预设与"
        "工程可切换接口；（2）完成向量检索与 RAG 问答，使识别结果可参与查询扩展；（3）在统一标注基准上"
        "给出可复现的定量对比与误差分析，明确各方法适用边界与推荐配置。",
    )

    add_heading_cn(doc, "1.4 研究内容与技术路线", 2)
    add_para(
        doc,
        "研究内容与技术路线调整为以识别为主线、检索与问答为支撑：（1）GO 数据组织、元数据与 FAISS 向量"
        "索引；（2）GO 术语识别：八类可区分管线及 Ensemble 三档预设、LLM 保守策略、RAG+LLM 候选规模等"
        "组合与消融；（3）语义检索与可选 IC/两阶段索引；（4）RAG 问答及识别驱动的查询扩展；（5）FastAPI"
        "服务化与 benchmark 脚本复现实验。路线概括为「数据与索引 → 多策略识别与对比 → 检索与 RAG →"
        "联评与总结」。",
    )

    add_heading_cn(doc, "1.5 论文组织结构", 2)
    add_para(
        doc,
        "第二章介绍 GO、向量检索、FAISS、图数据、RAG、术语识别及 IC 索引思想、指标形式化、服务化要点，"
        "并扩展水稻文本特点、可重复性规范及与主干课程的衔接；第三章以识别子系统为重点描述"
        "总体架构、各模块实现及数据与复现说明；第四章给出识别主实验、组合与消融、与问答联评及"
        "工程化评测与写作细节；第五章总结全文、展望并说明图表与答辩规范。正文不设附录。",
    )

    _thesis_exp.expand_chapter1_contributions(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter1_deep(doc, add_para, add_heading_cn)

    add_heading_cn(doc, "本章小结", 2)
    add_para(
        doc,
        "本章从机理文本到 GO 标识的映射需求出发，强调实体识别在全文工作中的枢纽地位，综述了问答、RAG 与"
        "本体识别相关研究，明确了以多路线识别与对比实验为重点的目标、内容与技术路线。",
    )

    doc.add_page_break()

    # ========== 第二章 ==========
    add_heading_cn(doc, "第二章 相关理论与技术基础", 1)

    add_heading_cn(doc, "2.1 Gene Ontology 与分子机理知识表示", 2)
    add_para(
        doc,
        "GO 通过有向无环图组织术语，并提供文本定义与同义词信息。基因产物可通过实验或计算证据注释到具体"
        "GO 节点。理解 BP/MF/CC 三类命名空间及其关系，有助于将自然语言中的机理描述映射到可计算的功能单元。",
    )

    add_heading_cn(doc, "2.2 文本向量化与语义检索", 2)
    add_para(
        doc,
        "将术语名称、定义及描述编码为稠密向量后，可使用余弦相似度等度量进行语义近邻搜索。生物医学领域"
        "预训练语言模型（如基于 PubMed 语料的句向量模型）在术语相似度任务上通常优于通用语料模型。",
    )

    add_heading_cn(doc, "2.3 近似最近邻与 FAISS 索引", 2)
    add_para(
        doc,
        "面对数万级以上的 GO 条目，暴力线性扫描代价较高。Facebook AI Similarity Search（FAISS）提供多种"
        "索引结构与量化策略，可在精度与速度之间折中。本文将向量与元数据分离存储，检索阶段返回 Top-K 候选"
        "及分数用于后续排序与提示构造。",
    )

    add_heading_cn(doc, "2.4 图数据库与知识组织", 2)
    add_para(
        doc,
        "Neo4j 等图数据库适合表达实体—关系网络。将 GO 术语及相关实体（如基因、文献节点等）以图模型存储，"
        "可在问答时补充结构化邻域信息，并与向量检索形成互补。",
    )

    add_heading_cn(doc, "2.5 大语言模型与检索增强生成", 2)
    add_para(
        doc,
        "大语言模型通过大规模自监督学习获得语言生成能力。检索增强生成将检索器得到的片段作为条件上下文，"
        "使生成「有依据」。在 GO 场景中，提示工程常结合 JSON 输出、角色设定，以及「仅从给定候选 GO 列表"
        "择项」等约束，以缩小输出空间、抑制与本体不一致的编造。",
    )

    add_heading_cn(doc, "2.6 生物医学场景下的术语识别与本体归一", 2)
    add_para(
        doc,
        "命名实体识别（NER）旨在从非结构化文本中定位实体边界与类型；本体归一（grounding / linking）则进一步"
        "将表面形式映射到受控词表的唯一标识符。对 GO 而言，输出通常为 GO:0005575 等形式，评价多采用集合级"
        "Precision / Recall / F1，并关注负例上的误报率。多策略融合（级联、加权投票、检索增强判别）可在"
        "高召回与低误报之间取得不同工作点，需在统一基准上对比方可客观选型。",
    )
    add_para(doc, "为便于说明，本研究将核心评价公式统一写为：Precision = TP / (TP + FP)，Recall = TP / (TP + FN)，F1 = 2PR / (P + R)。")
    add_para(doc, "若将每个句子的预测集合记为 Ŝ，金标准集合记为 S，则单句级匹配可写为 Jaccard = |Ŝ ∩ S| / |Ŝ ∪ S|，该值可用于辅助分析部分样例的重叠程度。")
    add_para(doc, "向量检索中可采用余弦相似度 sim(x, y) = (x · y) / (||x|| ||y||)，并以 Top-K 候选集合 C_K 进入后续判别或提示构造。")

    _thesis_exp.expand_chapter2(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter2_bioinformatics_rice(doc, add_para, add_heading_cn)

    add_heading_cn(doc, "2.15 本章小结", 2)
    add_para(
        doc,
        "本章在 GO 表示、向量检索、FAISS、图数据、RAG、术语识别归一、IC 与指标形式化、服务化部署以及"
        "水稻文本特点与可重复性规范等方面建立概念基础，为第三章多路线识别模块与第四章实验设计提供理论支撑。",
    )

    doc.add_page_break()

    # ========== 第三章 ==========
    add_heading_cn(doc, "第三章 系统总体设计与实现", 1)

    add_heading_cn(doc, "3.1 需求分析", 2)
    add_para(
        doc,
        "在功能上，除自然语言问句、带评分的语义检索与基于证据的 RAG 问答（含流式输出）外，需突出对"
        "GO 实体识别的完整支持：多种识别方法可经 REST 切换；集成模式支持多档预设；识别结果可回灌为"
        "查询扩展文本以影响检索。非功能需求包括：模块可替换（嵌入模型、索引路径、大模型 API）、"
        "脚本化 benchmark 可复现、延迟与误报可度量。",
    )

    add_heading_cn(doc, "3.2 总体架构", 2)
    add_para(
        doc,
        "数据层维护 go.obo、FAISS 索引与元数据、Neo4j 图谱等；服务层以 FastAPI 暴露 /ner、/ask 等接口；"
        "模型层包含句向量编码器与 Ollama 等兼容接口的大模型。逻辑上可将「识别子系统」视为相对独立的"
        "横切能力：其输出既可单独返回给调用方，也可在问答路径中与检索、RAG 串联，形成「识别 →（扩展）"
        "检索 → 生成」流水线。",
    )

    add_heading_cn(doc, "3.3 GO 术语识别子系统（本文工作重点）", 2)
    _thesis_exp.expand_chapter3_ner_algorithms(doc, add_para, add_heading_cn)

    add_heading_cn(doc, "3.4 数据与索引构建", 2)
    add_para(
        doc,
        "自 OBO 解析术语字段，构造用于编码的文本（名称、定义、同义词等），写入向量并维护与 FAISS 行号"
        "对齐的 metadata JSON。识别中的向量分支与全局问答检索可复用或分_profile 使用不同索引，以兼顾"
        "实验变量控制与部署灵活性。支持管理接口触发重建。",
    )

    add_heading_cn(doc, "3.5 检索模块", 2)
    add_para(
        doc,
        "对问句（或经识别扩展后的文本）编码，在 FAISS 上取 Top-K，返回带 node_type 等字段的元数据。"
        "可选 IC 压缩或两阶段粗精筛以调节延迟与命中质量，与识别实验可分开报告以免混淆变量。",
    )

    add_heading_cn(doc, "3.6 RAG 问答模块", 2)
    add_para(
        doc,
        "将检索命中格式化为证据块，拼装系统提示与用户问题，调用大模型生成答案，并结构化返回 sources。"
        "流式接口改善交互。提示中要求严格依据证据作答，避免虚构 GO 编号。",
    )

    add_heading_cn(doc, "3.7 识别驱动的查询扩展与接口", 2)
    add_para(
        doc,
        "问答路径可选在检索前调用 _expand_query_with_ner：将识别出的 GO_ID 与术语名拼接为 GO_HINTS 附加"
        "到查询，从而把识别误差对下游的影响纳入联评。对外提供 /health、/ner、/ask、/ask_stream、索引重建"
        "与 IC 索引切换等接口；配置经环境变量注入。",
    )

    add_heading_cn(doc, "3.8 本章小结", 2)
    add_para(
        doc,
        "本章除架构叙述外，重点说明了系统配置项、评测脚本、索引构建流程与批量测试流程之间的对应关系，"
        "保证实验描述可追溯至具体参数与执行步骤，而不在正文中直接展开具体代码文件名。",
    )

    _thesis_exp.expand_chapter3_post_data_repro(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter3_code_implementation(doc, add_para, add_heading_cn)

    doc.add_page_break()

    # ========== 第四章 ==========
    add_heading_cn(doc, "第四章 实验与结果分析", 1)

    add_heading_cn(doc, "4.1 实验环境", 2)
    add_para(
        doc,
        "实验硬件与操作系统环境按实际填写（如 CPU 型号、内存容量、是否使用 GPU、服务器或工作站系统版本等）。"
        "软件环境包括 Python 解释器版本、依赖库版本、Neo4j 版本、FAISS 构建参数、句向量模型名称与权重来源、"
        "Ollama 及所选用大模型名称与参数量等。",
    )

    add_heading_cn(doc, "4.2 数据集与评价指标", 2)
    add_para(
        doc,
        "识别主实验使用自建 GO 标注基准（如约 200 条规模、含 easy/medium/hard 及中英与混合句）。以句为单位"
        "统计预测 GO_ID 集合与金标准集合的 TP/FP/FN，汇总 micro 与 macro 的 Precision、Recall、F1，并记录"
        "负例句上的误报率（NegFP%）、全对率（exact match）及单句平均耗时。RAG+LLM 分支可通过统一的批量"
        "评测流程输出 micro/macro 与索引配置。检索与问答联评可采用 Recall@K、人工相关性打分或 API 批量脚本"
        "对比有无 GO_HINTS 扩展时的 sources 差异。",
    )
    add_heading_cn(doc, "4.2.1 指标表标题与说明", 3)
    add_caption(doc, "表 4-1  GO 术语识别主实验总体指标对比（Precision / Recall / F1 / NegFP% / AvgTime）")
    add_para(
        doc,
        "该表建议置于本节正文后，按方法名称逐行列出 Dict、NLTK、Vector、LLM、LLM（保守）、Ensemble 等方法的总体结果，并在表下注明数据来源为统一基准批量评测输出。",
    )
    add_caption(doc, "表 4-2  GO 术语识别按语言分层结果对比")
    add_para(
        doc,
        "建议展示 zh、en、mixed 三个子集上的宏平均 F1，便于说明不同方法对中文、英文与混合句的适应性差异。",
    )
    add_caption(doc, "表 4-3  GO 术语识别按难度分层结果对比")
    add_para(
        doc,
        "建议展示 easy、medium、hard 与 hard-negative 子集上的 P/R/F1 及误报情况，用于解释保守策略与集成策略在困难样本上的折中。",
    )
    add_caption(doc, "图 4-1  数据集构成与评测流程示意图")
    add_para(
        doc,
        "图中可依次展示输入文本、候选生成、识别预测、指标统计与结果汇总的闭环流程，作为本章评测设计的总览图。",
    )
    add_caption(doc, "图 4-2  GO 术语识别主实验结果对比图")
    add_para(
        doc,
        "可用柱状图或雷达图展示各方法的 F1 与 NegFP% 折中，建议将 LLM（保守）与 Ensemble 作为重点对照对象。",
    )
    add_para(
        doc,
        "插图位置说明：上述表格与图示可在本节段落后直接插入；若 Word 排版时出现跨页，可优先保证表 4-1 紧跟本节首段，图 4-1 放在指标定义之后，图 4-2 放在结果分析之前。",
    )

    add_heading_cn(doc, "4.3 GO 术语识别：多方法横向对比", 2)
    add_caption(doc, "表 4-4  GO 术语识别多方法横向对比结果")
    add_para(
        doc,
        "本节正文可据此展开总体趋势分析。在相同基准与实现版本下，对 Dict、NLTK、Vector、LLM（非保守）、LLM（保守）、Ensemble 等配置跑通全流程并制表。根据当前汇总结果：NLTK 的 macro F1 约为 0.437，单句平均耗时极低；原始 LLM 的 macro F1 约为 0.480，但 NegFP% 约 28.6%；LLM（保守模式）macro F1 约 0.548，NegFP% 降至约 5.7%，在综合指标与误报控制上优于若干基线；Dict 与 Vector 在综合 F1 上相对较低但在部分子集上仍具对照价值。Ensemble 曾出现极高负例误报，提示融合阈值需与数据分布联合调参，不宜直接作为默认上线配置。",
    )
    add_caption(doc, "图 4-3  各方法 Precision、Recall、F1 与 NegFP% 对比图")
    add_para(
        doc,
        "建议采用分组柱状图，将精度与误报率并列呈现，便于直观看出保守 LLM 在 F1 提升和误报下降上的平衡优势。",
    )
    add_para(
        doc,
        "插图位置说明：图 4-3 可置于本节结果说明段之后、讨论段之前；若附带更多子图，可采用 4-3(a)、4-3(b) 的形式分别展示总体指标与负例误报情况。",
    )

    add_heading_cn(doc, "4.4 组合、预设与消融实验", 2)
    add_caption(doc, "表 4-5  Ensemble 三档预设与关键超参消融结果表")
    add_para(
        doc,
        "建议对 strict / balanced / recall 三种配置分别统计 P/R/F1 与 NegFP%，并补充 candidate_top_k、vector_threshold 等关键参数。纵向实验：对同一方法调整关键超参（如 LLM 是否保守、RAG+LLM 的 candidate_top_k、向量相似度阈值等），观察 P/R/F1 与 NegFP% 的变化。横向实验：在固定超参下对比八类管线中的子集（以实际已跑通的配置为准）。对 Ensemble，分别报告 strict / balanced / recall 三档预设的识别结果与误报特征。消融项示例：去除检索增强中的词典或向量一路候选；去除查询扩展环节，比较下游检索命中与问答忠实度。",
    )
    add_caption(doc, "图 4-4  Ensemble 预设与消融对比图")
    add_para(
        doc,
        "可使用折线图或雷达图展示不同预设在召回与误报上的变化。",
    )
    add_para(
        doc,
        "插图位置说明：图 4-4 建议放在消融结果总结段之后，以便读者先看到文字解释，再看到图形化趋势。",
    )
    add_para(
        doc,
        "图 4-4  Ensemble 预设与消融对比图；可使用折线图或雷达图展示不同预设在召回与误报上的变化。",
    )
    add_para(
        doc,
        "插图位置说明：图 4-4 建议放在消融结果总结段之后，以便读者先看到文字解释，再看到图形化趋势。",
    )

    add_heading_cn(doc, "4.5 检索、RAG 与联评（辅助）", 2)
    add_para(
        doc,
        "表 4-6  检索与 RAG 联评结果表；建议展示 Top-K 召回率、sources 数量、回答一致性与查询扩展前后"
        "的变化。"
        "在识别结果稳定的前提下，报告语义检索 Top-K、可选 IC 索引对召回的影响，以及开启/关闭 NER 查询扩展"
        "时典型问句的 sources 与回答差异。该部分篇幅可小于识别主实验，但用于说明全系统闭环。",
    )
    add_para(
        doc,
        "图 4-5  识别—检索—生成联动流程图；可展示输入问题、识别扩展、证据召回与答案生成之间的关系。",
    )

    add_heading_cn(doc, "4.6 定性结果与典型案例", 2)
    add_para(
        doc,
        "选取中英混合句、长句多 GO、以及 hard-negative 样例，截图或摘录 JSON：展示各方法预测集合与金标准"
        "的差异，辅以简短语言学分析（同义、省略、领域隐喻导致的漏检等）。",
    )

    add_heading_cn(doc, "4.7 结果讨论", 2)
    add_para(
        doc,
        "归纳各方法适用场景：词典适合标准术语；NLTK 适合字面匹配强的英文；向量缓解表述变体但易引入"
        "语义近邻误配；保守 LLM 与 RAG+LLM 在控制幻觉与约束输出空间上的利弊；集成预设对召回—误报折中"
        "的敏感性。说明中文子集上词典/NLTK 可能偏弱、需依赖大模型分支的实验现象及改进思路。",
    )

    _thesis_exp.expand_chapter4(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter4_engineering(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter4_measured_results(doc, add_para, add_heading_cn)

    add_heading_cn(doc, "4.18 本章小结", 2)
    add_para(
        doc,
        "本章除数据集与指标外，将整体结果、按语言结果、按难度结果与保守提示、默认集成构造对照说明，"
        "并交代逐样本 JSON、批量评测输出、IC 消融与 API 联评、向量存储与 RAG 调用参数的引用方式；"
        "定稿时若重新跑分，应以新版 JSON 替换表中数字。实验数据应以最新产物为准。",
    )

    doc.add_page_break()

    # ========== 第五章 ==========
    add_heading_cn(doc, "第五章 总结与展望", 1)

    add_heading_cn(doc, "5.1 本文工作总结", 2)
    add_para(
        doc,
        "（1）设计并实现覆盖八类路线的 GO 术语识别子系统，含级联、检索增强判别与多预设加权集成，并完成"
        "与统一基准对齐的批量评测；（2）搭建 FAISS 向量检索与可选 IC/两阶段索引，支撑候选生成与 RAG；"
        "（3）实现 RAG 问答、流式输出、来源追溯及识别驱动的查询扩展，形成完整应用链路。",
    )

    add_heading_cn(doc, "5.2 不足与展望", 2)
    add_para(
        doc,
        "不足包括：基准规模与领域覆盖面仍可扩大；中文与同义词资源不足制约词典与句法分支；大模型与"
        "多路调用带来延迟与成本；集成策略对负例敏感需持续调参。展望：引入更多水稻专用库做联合 grounding；"
        "半自动标注与主动学习迭代基准；轻量小模型承担粗筛、大模型仅做歧义消解以降本；与文献全文、"
        "多跳推理结合提升可解释性。",
    )

    _thesis_exp.expand_chapter5(doc, add_para, add_heading_cn)
    _thesis_exp.expand_chapter5_submission(doc, add_para, add_heading_cn)

    doc.add_page_break()

    # ========== 参考文献 ==========
    add_heading_cn(doc, "参考文献", 1)
    refs = [
        "Ashburner M, et al. Gene ontology: tool for the unification of biology[J]. Nature Genetics, 2000, 25(1): 25-29.",
        "The Gene Ontology Consortium. The Gene Ontology resource: enriching a GOld mine[J]. Nucleic Acids Research, 2021, 49(D1): D325-D334.",
        "Lewis N E, et al. Rice pathways: databases and tools for rice metabolic and regulatory networks[J]. Plant Physiology, 2015, 169(3): 1516-1524.",
        "Devlin J, et al. BERT: Pre-training of deep bidirectional transformers for language understanding[C]//NAACL-HLT. 2019.",
        "Reimers N, Gurevych I. Sentence-BERT: Sentence embeddings using Siamese BERT-networks[C]//EMNLP-IJCNLP. 2019.",
        "Chalkidis I, et al. An exploration of neural architectures for taxonomy learning[C]//EMNLP Workshop. 2020.",
        "Johnson J, Douze M, Jégou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2019, 7(3): 535-547.",
        "Lewis P, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//NeurIPS. 2020.",
        "Guu K, et al. REALM: Retrieval-augmented language model pre-training[J]. arXiv preprint arXiv:2002.08909, 2020.",
        "Izacard G, Grave E. Leveraging passage retrieval with generative models for open domain question answering[C]//EACL. 2021.",
        "Gao Y, et al. Retrieval-augmented generation: A survey[J]. arXiv preprint arXiv:2312.10997, 2023.",
        "Wei J, et al. Chain-of-thought prompting elicits reasoning in large language models[C]//NeurIPS. 2022.",
        "Brown T, et al. Language models are few-shot learners[C]//NeurIPS. 2020.",
        "Touvron H, et al. LLaMA: Open and efficient foundation language models[J]. arXiv preprint arXiv:2302.13971, 2023.",
        "Neo4j Inc. Neo4j Graph Database Documentation[EB/OL]. https://neo4j.com/docs/, 2024.",
        "FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2024.",
        "Robinson I, Webber J, Eifrem E. Graph Databases: New Opportunities for Connected Data[M]. O'Reilly Media, 2015.",
        "Jurafsky D, Martin J H. Speech and Language Processing: An Introduction to Natural Language Processing[M]. Pearson, 2024.",
        "王斌, 李锦涛. 知识图谱：方法、实践与应用[M]. 北京: 电子工业出版社, 2019.",
        "刘知远, 孙茂松. 知识表示学习研究进展[J]. 计算机研究与发展, 2016, 53(2): 247-261.",
        "张奇, 桂韬, 黄萱菁. 大规模语言模型：从理论到实践[M]. 北京: 电子工业出版社, 2024.",
        "Lample M, et al. Neural architectures for named entity recognition[C]//NAACL-HLT. 2016.",
        "国家标准化管理委员会. GB/T 7714-2015 信息与文献 参考文献著录规则[S]. 北京: 中国标准出版社, 2015.",
    ]
    for i, line in enumerate(refs, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"[{i}] {line}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        set_run_eastasia_font(run, "宋体", 10.5)

    doc.add_page_break()

    # 致谢
    add_heading_cn(doc, "致  谢", 1)
    add_para(
        doc,
        "本论文是在导师的悉心指导下完成的。导师在选题、实体识别方案论证与论文结构上给予了耐心指导，使本人"
        "在生物信息应用与工程评测方面获益良多。感谢学院各位老师的教导，感谢同学在批量 benchmark 与"
        "服务器环境调试中的协助，感谢家人的支持。论文中错漏之处恳请批评指正。",
    )

    doc.add_page_break()

    return doc


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(
        out_dir,
        "毕业论文_基于语言大模型的水稻分子机理数据检索和知识问答.docx",
    )
    # 同步生成纯英文文件名副本，方便在资源管理器中按字母排序查找、避免部分环境对长中文路径显示异常
    alt_path = os.path.join(out_dir, "graduation_thesis_rice_go_rag.docx")

    doc = build_document()
    abs_main = os.path.abspath(out_path)
    abs_alt = os.path.abspath(alt_path)
    txt_path = os.path.join(out_dir, "毕业论文_基于语言大模型的水稻分子机理数据检索和知识问答.txt")
    txt_alt_path = os.path.join(out_dir, "graduation_thesis_rice_go_rag.txt")
    try:
        doc.save(out_path)
        shutil.copy2(out_path, alt_path)
        plain_text = document_to_plain_text(doc)
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(plain_text)
        with open(txt_alt_path, "w", encoding="utf-8") as f:
            f.write(plain_text)
        print("已生成（与脚本同目录）:")
        print(" ", abs_main)
        print(" ", abs_alt)
        print(" ", os.path.abspath(txt_path))
        print(" ", os.path.abspath(txt_alt_path))
    except PermissionError:
        # 常见原因：同名 docx 正在被 Word 打开，无法覆盖
        spare_zh = os.path.join(out_dir, "毕业论文_基于语言大模型的水稻分子机理数据检索和知识问答_生成稿.docx")
        spare_en = os.path.join(out_dir, "graduation_thesis_rice_go_rag_generated.docx")
        spare_txt = os.path.join(out_dir, "毕业论文_基于语言大模型的水稻分子机理数据检索和知识问答_生成稿.txt")
        spare_txt_en = os.path.join(out_dir, "graduation_thesis_rice_go_rag_generated.txt")
        doc.save(spare_zh)
        shutil.copy2(spare_zh, spare_en)
        plain_text = document_to_plain_text(doc)
        with open(spare_txt, "w", encoding="utf-8") as f:
            f.write(plain_text)
        with open(spare_txt_en, "w", encoding="utf-8") as f:
            f.write(plain_text)
        print("原文件名被占用（请先关闭 Word 再运行以覆盖原文件）。已改为写入:")
        print(" ", os.path.abspath(spare_zh))
        print(" ", os.path.abspath(spare_en))
        print(" ", os.path.abspath(spare_txt))
        print(" ", os.path.abspath(spare_txt_en))


if __name__ == "__main__":
    main()
